"""
Redline helper for the /legal skill.

Wraps python-docx with utilities for:
  - Striking through arbitrary text inside a paragraph (red + strikethrough)
  - Inserting red text without strikethrough
  - Replacing a paragraph entirely
  - Adding new sections after an anchor paragraph
  - Verifying edits landed (post-save sanity check)

Handles two common silent-failure modes:
  1. Curly apostrophes (U+2019) in source contracts vs. straight apostrophes in code
  2. Multi-run text (a single paragraph may span many runs at formatting boundaries)

Usage:
    from redline_helper import RedlineDoc

    doc = RedlineDoc('/path/to/source.docx')
    doc.strike_substring(doc.paragraphs[5], "any prior agreements between the parties")
    doc.replace_paragraph(13, "Old text...", "New replacement text...")
    doc.add_section_after(27, "28. New Section. ", "Body of the new section...")
    doc.save('/path/to/output.docx')
    doc.verify([
        ("MFN strike", "less favorable than the lowest price", "strike"),
        ("New §28 inserted", "New Section", "red"),
    ])
"""

from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

RED = RGBColor(0xC0, 0x00, 0x00)


class RedlineDoc:
    def __init__(self, source_path: str):
        self.doc = Document(source_path)
        self.source_path = source_path
        self.output_path = None

    # ----- properties -----
    @property
    def paragraphs(self):
        return self.doc.paragraphs

    # ----- formatting primitives -----
    @staticmethod
    def _red_format(run, *, strike=False, bold=False, italic=False):
        run.font.color.rgb = RED
        if strike:
            run.font.strike = True
        if bold:
            run.font.bold = True
        if italic:
            run.font.italic = True

    def add_red_run(self, paragraph, text, *, bold=False, italic=False):
        r = paragraph.add_run(text)
        self._red_format(r, bold=bold, italic=italic)
        return r

    # ----- whole-paragraph operations -----
    def strike_paragraph(self, paragraph):
        """Strike through every run in the paragraph and color it red."""
        for r in paragraph.runs:
            r.font.color.rgb = RED
            r.font.strike = True

    def insert_paragraph_after(self, anchor_paragraph, text=None, *, bold=False, italic=False):
        """Insert a new <w:p> immediately after the anchor and return as Paragraph."""
        new_p = OxmlElement("w:p")
        anchor_paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, anchor_paragraph._parent)
        if text:
            self.add_red_run(new_para, text, bold=bold, italic=italic)
        return new_para

    def add_section_after(self, anchor_paragraph_or_idx, header, body, *, header_bold=True):
        """Add a new section after an anchor: bold red header + body text."""
        if isinstance(anchor_paragraph_or_idx, int):
            anchor = self.paragraphs[anchor_paragraph_or_idx]
        else:
            anchor = anchor_paragraph_or_idx
        new_para = self.insert_paragraph_after(anchor)
        if header:
            self.add_red_run(new_para, header, bold=header_bold)
        if body:
            self.add_red_run(new_para, body)
        return new_para

    # ----- substring strike / replace (handles multi-run text + curly quotes) -----
    def _candidates(self, text):
        """Return the text plus a curly-apostrophe variant for resilient matching."""
        variants = [text]
        if "'" in text:
            variants.append(text.replace("'", "’"))
        if "’" in text:
            variants.append(text.replace("’", "'"))
        return variants

    def strike_substring(self, paragraph, substring):
        """
        Strike a substring within a paragraph by rebuilding all runs.
        Returns True if found and struck, False otherwise.
        Tries both straight and curly apostrophes.
        """
        for candidate in self._candidates(substring):
            full = paragraph.text
            if candidate not in full:
                continue
            pre, _, post = full.partition(candidate)
            for r in list(paragraph.runs):
                r._r.getparent().remove(r._r)
            if pre:
                paragraph.add_run(pre)
            sr = paragraph.add_run(candidate)
            self._red_format(sr, strike=True)
            if post:
                paragraph.add_run(post)
            return True
        return False

    def replace_paragraph(self, paragraph_or_idx, old_text, new_red_text):
        """
        Strike old_text and insert new_red_text in its place within the paragraph.
        new_red_text appears as red (insertion).
        """
        if isinstance(paragraph_or_idx, int):
            paragraph = self.paragraphs[paragraph_or_idx]
        else:
            paragraph = paragraph_or_idx
        for candidate in self._candidates(old_text):
            full = paragraph.text
            if candidate not in full:
                continue
            pre, _, post = full.partition(candidate)
            for r in list(paragraph.runs):
                r._r.getparent().remove(r._r)
            if pre:
                paragraph.add_run(pre)
            sr = paragraph.add_run(candidate)
            self._red_format(sr, strike=True)
            if new_red_text:
                nr = paragraph.add_run(new_red_text)
                self._red_format(nr)
            if post:
                paragraph.add_run(post)
            return True
        return False

    def strike_phrases_preserving_runs(self, paragraph, phrases):
        """
        Strike multiple phrases within a paragraph while preserving the
        formatting of all surrounding text. Use this when a paragraph
        already contains red insertions you don't want to clobber.

        Returns the count of phrases successfully struck.
        """
        # Capture run-level data
        existing = []
        for r in paragraph.runs:
            existing.append({
                'text': r.text,
                'color': r.font.color.rgb,
                'strike': r.font.strike or False,
                'italic': r.font.italic or False,
                'bold': r.font.bold or False,
            })
        # Build per-character format array
        char_format = []
        for e in existing:
            for c in e['text']:
                char_format.append({
                    'char': c, 'color': e['color'], 'strike': e['strike'],
                    'italic': e['italic'], 'bold': e['bold'],
                })
        full = ''.join(c['char'] for c in char_format)

        struck = 0
        for phrase in phrases:
            for candidate in self._candidates(phrase):
                idx = full.find(candidate)
                if idx == -1:
                    continue
                for k in range(idx, idx + len(candidate)):
                    char_format[k]['strike'] = True
                    char_format[k]['color'] = RED
                struck += 1
                break

        if struck == 0:
            return 0

        # Rebuild runs by grouping consecutive chars with identical formatting
        for r in list(paragraph.runs):
            r._r.getparent().remove(r._r)
        if not char_format:
            return struck

        cur_text = char_format[0]['char']
        cur_fmt = (char_format[0]['color'], char_format[0]['strike'],
                   char_format[0]['italic'], char_format[0]['bold'])
        for cf in char_format[1:]:
            fmt = (cf['color'], cf['strike'], cf['italic'], cf['bold'])
            if fmt == cur_fmt:
                cur_text += cf['char']
            else:
                self._add_formatted_run(paragraph, cur_text, cur_fmt)
                cur_text = cf['char']
                cur_fmt = fmt
        self._add_formatted_run(paragraph, cur_text, cur_fmt)
        return struck

    @staticmethod
    def _add_formatted_run(paragraph, text, fmt):
        color, strike, italic, bold = fmt
        nr = paragraph.add_run(text)
        if color:
            nr.font.color.rgb = color
        if strike:
            nr.font.strike = True
        if italic:
            nr.font.italic = True
        if bold:
            nr.font.bold = True

    # ----- save / verify -----
    def save(self, output_path: str):
        self.doc.save(output_path)
        self.output_path = output_path

    def verify(self, checks):
        """
        Re-open the saved file and verify each (label, snippet, mode) check.
        mode ∈ {"strike", "red", "absent"}.
        Prints OK/MISS for each, returns list of (label, ok) tuples.
        """
        if not self.output_path:
            raise RuntimeError("Call save() before verify()")
        doc = Document(self.output_path)
        results = []
        for label, snippet, mode in checks:
            found = False
            for p in doc.paragraphs:
                for r in p.runs:
                    if mode == "strike" and snippet in r.text and r.font.strike:
                        found = True
                        break
                    if mode == "red" and snippet in r.text and r.font.color.rgb == RED and not r.font.strike:
                        found = True
                        break
                    if mode == "absent" and snippet in r.text:
                        found = True
                        break
                if found:
                    break
            ok = (not found) if mode == "absent" else found
            results.append((label, ok))
            tag = "OK  " if ok else "MISS"
            print(f"{tag} {label}")
        return results
